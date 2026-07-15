"""
Resume & JD text parser.
Supports PDF (digital) and DOCX formats.
Also detects resume sections for structured feedback.
"""

import io
import re
from typing import Dict, Optional

from pdfminer.high_level import extract_text as pdf_extract_text
from pdfminer.pdfparser import PDFSyntaxError
import docx


SECTION_HEADERS = {
    "summary": [
        r"\bsummary\b", r"\bprofile\b", r"\bobjective\b",
        r"\babout\b", r"\bprofessional summary\b",
    ],
    "skills": [
        r"\bskills\b", r"\btechnical skills\b", r"\bcore competencies\b",
        r"\bkey skills\b", r"\btechnologies\b", r"\btools\b",
    ],
    "experience": [
        r"\bexperience\b", r"\bwork experience\b", r"\bprofessional experience\b",
        r"\bemployment\b", r"\bwork history\b", r"\bcareer\b",
    ],
    "education": [
        r"\beducation\b", r"\bacademic background\b", r"\bqualifications\b",
        r"\bdegrees?\b",
    ],
    "projects": [
        r"\bprojects?\b", r"\bpersonal projects?\b", r"\bside projects?\b",
        r"\bportfolio\b",
    ],
    "certifications": [
        r"\bcertifications?\b", r"\bcertificates?\b", r"\blicenses?\b",
        r"\baccreditations?\b",
    ],
}


def parse_pdf(file_bytes: bytes) -> str:
    """Extract text from a digital PDF."""
    try:
        text = pdf_extract_text(io.BytesIO(file_bytes))
        return _clean_text(text)
    except PDFSyntaxError:
        raise ValueError("Could not parse PDF. Make sure it's a digital (not scanned) PDF.")
    except Exception as e:
        raise ValueError(f"PDF parsing error: {str(e)}")


def parse_docx(file_bytes: bytes) -> str:
    """Extract text from a DOCX file."""
    try:
        doc = docx.Document(io.BytesIO(file_bytes))
        paragraphs = [para.text for para in doc.paragraphs]
        # Also get text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    paragraphs.append(cell.text)
        text = "\n".join(p for p in paragraphs if p.strip())
        return _clean_text(text)
    except Exception as e:
        raise ValueError(f"DOCX parsing error: {str(e)}")


def parse_resume(file_bytes: bytes, filename: str) -> str:
    """Auto-detect format and parse resume."""
    filename_lower = filename.lower()
    if filename_lower.endswith(".pdf"):
        return parse_pdf(file_bytes)
    elif filename_lower.endswith((".docx", ".doc")):
        return parse_docx(file_bytes)
    else:
        raise ValueError(f"Unsupported file format: {filename}. Only PDF and DOCX are supported.")


def detect_sections(resume_text: str) -> Dict[str, str]:
    """
    Split resume text into named sections.
    Returns a dict like {'skills': '...', 'experience': '...'}.
    """
    lines = resume_text.split("\n")
    sections: Dict[str, str] = {}
    current_section: Optional[str] = None
    current_content: list[str] = []

    for line in lines:
        stripped = line.strip()
        if not stripped:
            if current_content:
                current_content.append("")
            continue

        detected = _detect_section_header(stripped)
        if detected:
            # Save previous section
            if current_section and current_content:
                sections[current_section] = "\n".join(current_content).strip()
            current_section = detected
            current_content = []
        else:
            current_content.append(stripped)

    # Save last section
    if current_section and current_content:
        sections[current_section] = "\n".join(current_content).strip()

    # If no sections detected, put everything under 'general'
    if not sections:
        sections["general"] = resume_text

    return sections


def _detect_section_header(line: str) -> Optional[str]:
    """Check if a line is a known section header."""
    # Section headers are usually short (< 50 chars) and may be ALL CAPS
    if len(line) > 60:
        return None

    line_lower = line.lower().strip()
    # Strip common punctuation
    line_clean = re.sub(r"[:\-–—_*#]", "", line_lower).strip()

    for section_name, patterns in SECTION_HEADERS.items():
        for pattern in patterns:
            if re.search(pattern, line_clean):
                return section_name
    return None


def _clean_text(text: str) -> str:
    """Clean extracted text: normalize whitespace, remove junk characters."""
    if not text:
        return ""
    # Normalize whitespace
    text = re.sub(r"\r\n", "\n", text)
    text = re.sub(r"\r", "\n", text)
    # Remove multiple consecutive blank lines
    text = re.sub(r"\n{3,}", "\n\n", text)
    # Remove non-printable characters (except newlines)
    text = re.sub(r"[^\x20-\x7E\n]", " ", text)
    return text.strip()


def get_resume_stats(resume_text: str) -> Dict:
    """Quick structural stats used in format quality scoring."""
    lines = [l.strip() for l in resume_text.split("\n") if l.strip()]
    word_count = len(resume_text.split())
    bullet_count = sum(1 for l in lines if l.startswith(("-", "•", "*", "·")))
    sections = detect_sections(resume_text)
    has_contact = bool(
        re.search(r"[\w.+-]+@[\w-]+\.[a-z]{2,}", resume_text, re.IGNORECASE)
    )
    return {
        "word_count": word_count,
        "line_count": len(lines),
        "bullet_count": bullet_count,
        "section_count": len(sections),
        "sections_found": list(sections.keys()),
        "has_contact_info": has_contact,
    }
